"""Post-process a pandoc-generated .docx to apply professional formatting.

Usage:
    python scripts/format_report.py <path_to_report.docx>

Applies:
- Table borders (all cells, dark gray)
- Header row: dark blue background (#003366) with bold white text
- Alternating row shading (light gray #F2F2F2 on even data rows)
- Consistent cell font (Calibri 10pt)
- Paragraph spacing: space before/after headings and between body paragraphs
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy


HEADER_BG = "003366"
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW_BG = "F2F2F2"
BORDER_COLOR = "999999"
FONT_NAME = "Calibri"
FONT_SIZE = Pt(10)
HEADER_FONT_SIZE = Pt(10)


def set_cell_borders(cell, color=BORDER_COLOR, size="4"):
    """Set all four borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f"</w:tcBorders>"
    )

    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(borders)


def set_cell_shading(cell, color):
    """Set background shading on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shading)


def format_cell_text(cell, bold=False, color=None, font_name=FONT_NAME, font_size=FONT_SIZE):
    """Format all runs in a cell's paragraphs."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = bold
            if color:
                run.font.color.rgb = color


def set_table_borders(table, color=BORDER_COLOR, size="4"):
    """Set inside borders on the table element itself."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f"</w:tblBorders>"
    )

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def format_table(table):
    """Apply full formatting to a single table."""
    # Set table-level borders
    set_table_borders(table)

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_borders(cell)

            if row_idx == 0:
                # Header row
                set_cell_shading(cell, HEADER_BG)
                format_cell_text(cell, bold=True, color=HEADER_FG, font_size=HEADER_FONT_SIZE)
            else:
                # Data rows - alternating shading
                if row_idx % 2 == 0:
                    set_cell_shading(cell, ALT_ROW_BG)
                format_cell_text(cell, bold=False, color=RGBColor(0x1A, 0x1A, 0x1A))


def format_paragraphs(doc):
    """Apply spacing to headings and body paragraphs."""
    # Spacing config: (space_before_pt, space_after_pt, line_spacing)
    HEADING_SPACING = {
        "Heading 1": (24, 12, 1.15),
        "Heading 2": (18, 6, 1.15),
        "Heading 3": (12, 6, 1.15),
    }
    BODY_SPACE_AFTER = Pt(8)
    BODY_LINE_SPACING = 1.15

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name

        if style_name in HEADING_SPACING:
            before, after, line_sp = HEADING_SPACING[style_name]
            pf = paragraph.paragraph_format
            pf.space_before = Pt(before)
            pf.space_after = Pt(after)
            pf.line_spacing = line_sp

        elif style_name in ("Normal", "Body Text", "First Paragraph"):
            pf = paragraph.paragraph_format
            pf.space_after = BODY_SPACE_AFTER
            pf.line_spacing = BODY_LINE_SPACING


def main():
    if len(sys.argv) < 2:
        print("Usage: python format_report.py <report.docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    doc = Document(docx_path)

    # Format all tables
    table_count = 0
    for table in doc.tables:
        format_table(table)
        table_count += 1

    # Format paragraph spacing
    format_paragraphs(doc)

    # Determine output path
    if len(sys.argv) >= 3:
        out_path = sys.argv[2]
    else:
        out_path = docx_path

    # Try saving; if file is locked, try _formatted suffix
    try:
        doc.save(out_path)
        print(f"Formatted {table_count} tables + paragraph spacing in {out_path}")
    except PermissionError:
        alt_path = out_path.replace(".docx", "_formatted.docx")
        doc.save(alt_path)
        print(f"Formatted {table_count} tables + paragraph spacing -> {alt_path} (original file was locked)")


if __name__ == "__main__":
    main()
